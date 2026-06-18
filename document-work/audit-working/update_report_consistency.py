from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "document-work" / "FYP Report_Rabih hajj hassan_Latest.docx"
OUTPUT = ROOT / "document-work" / "FYP Report_Rabih hajj hassan_Latest_consistency_audited.docx"


def replace_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def set_paragraph(doc, index, text, style=None):
    paragraph = doc.paragraphs[index]
    replace_paragraph_text(paragraph, text)
    if style:
        paragraph.style = style
    return paragraph


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    else:
        new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def delete_table(table):
    element = table._tbl
    element.getparent().remove(element)


def insert_table_after(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6))
    paragraph._p.addnext(table._tbl)
    return table


def mark_row_as_table_header(row):
    tr_props = row._tr.get_or_add_trPr()
    if tr_props.find(qn("w:tblHeader")) is not None:
        return
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_props.append(header)


def normalize_heading_styles(doc):
    if "Heading4" in [style.name for style in doc.styles]:
        doc.styles["Heading4"].name = "Heading 4"

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading3":
            paragraph.style = "Heading 3"
        if paragraph.style.name.startswith("Heading") and not paragraph.text.strip():
            paragraph.style = "Normal"


def clone_run_format(source_paragraph, target_paragraph):
    if not source_paragraph.runs or not target_paragraph.runs:
        return
    source_run_props = source_paragraph.runs[0]._r.rPr
    if source_run_props is None:
        return
    target_run = target_paragraph.runs[0]
    if target_run._r.rPr is not None:
        target_run._r.remove(target_run._r.rPr)
    target_run._r.insert(0, deepcopy(source_run_props))


def main():
    doc = Document(SOURCE)
    functional_requirements_anchor = doc.paragraphs[232]
    architecture_layer_anchor = doc.paragraphs[333]
    references_anchor = doc.paragraphs[528]
    acronyms_anchor = doc.paragraphs[537]
    table_list_anchor = doc.paragraphs[571]
    appendix_structure_anchor = doc.paragraphs[631]
    risk_heading_anchor = doc.paragraphs[193]
    toc_deployment_anchor = doc.paragraphs[126]
    deployment_body_anchor = doc.paragraphs[410]
    toc_paragraphs_to_delete = [doc.paragraphs[index] for index in range(127, 134)]

    # Front matter and visible TOC cleanup.
    set_paragraph(doc, 17, "Spring 2026")
    set_paragraph(doc, 18, "")
    set_paragraph(doc, 44, "")
    set_paragraph(doc, 80, "CHAPTER II: PROJECT REQUIREMENTS AND CONSTRAINTS 4")
    set_paragraph(doc, 78, "4. Risk Management and Scope Control 3")
    set_paragraph(doc, 79, "5. Report Outline 3")
    set_paragraph(doc, 83, "2.1. Team management 4")
    set_paragraph(doc, 96, "CHAPTER III: EXISTING SOLUTIONS 9")
    set_paragraph(doc, 100, "3.1.\tEstablished reporting platform patterns\t10")
    set_paragraph(doc, 104, "3.2.\tCommunity-centered reporting patterns\t11")
    set_paragraph(doc, 121, "2.1. Input block 19")
    set_paragraph(doc, 122, "2.2. Processing block 19")
    set_paragraph(doc, 123, "2.3. Communication block 19")

    # Chapter I/II cleanup.
    set_paragraph(doc, 186, (
        "The objective of this project is to develop a functional mobile-based application that enables users to submit "
        "and monitor safety-related reports using geolocation data, images, videos, and textual descriptions. The current "
        "implementation also includes a web-based moderator dashboard, a law-enforcement workspace, role-based administration, "
        "witness corroboration, saved-area and notification features, and machine-learning support for classification, risk "
        "analysis, duplicate detection, media evidence judgment, and analytics insights. The project is relevant to computer "
        "science students, civic technology initiatives, and organizations seeking practical applications of software "
        "engineering and intelligent systems in public safety."
    ))
    set_paragraph(doc, 188, (
        "The methodology adopted in this project follows an engineering design approach, emphasizing iterative development, "
        "modular system architecture, and continuous testing. The system is developed as a cross-platform mobile application "
        "using React Native and Expo, supported by a Node.js and Express backend, a React and Vite moderator dashboard, a "
        "PostgreSQL/PostGIS database, Socket.IO realtime communication, and a Python FastAPI machine-learning service."
    ))
    set_paragraph(doc, 190, (
        "The development process began with identifying functional requirements, including user authentication, incident "
        "reporting, geolocation tracking, photo and video upload support, offline report queueing, real-time notifications, "
        "witness corroboration, report categorization, duplicate detection, advisory media judgment, moderation workflows, "
        "law-enforcement follow-up, analytics, and administrative tools. Based on these requirements, the project was "
        "organized into independent modules: mobile application, moderator dashboard, backend API services, machine-learning "
        "service, database layer, and shared constants."
    ))
    set_paragraph(doc, 193, "Risk Management and Scope Control")
    set_paragraph(doc, 199, "CHAPTER II: PROJECT REQUIREMENTS AND CONSTRAINTS")
    set_paragraph(doc, 205, "Team management")
    set_paragraph(doc, 206, "")
    set_paragraph(doc, 211, (
        "Time management was a critical factor due to the complexity of the system and the academic constraints imposed by "
        "the semester schedule. The project was divided into several phases, including research, system design, mobile "
        "application development, backend implementation, dashboard implementation, machine-learning integration, testing, "
        "and documentation. Each phase was assigned a specific timeframe to prevent scope creep and ensure steady progress."
    ), style="Normal")
    set_paragraph(doc, 213, (
        "A Gantt chart was used to visualize the project timeline and monitor progress across different development stages. "
        "The chart outlined key milestones such as technology familiarization, system design, mobile application development, "
        "backend implementation, moderator dashboard development, machine-learning integration, system testing, and report "
        "writing."
    ), style="Normal")
    set_paragraph(doc, 217, (
        "The project was developed with minimal financial cost. Most development tools and frameworks used, including Expo, "
        "React Native, React, Vite, Node.js, Python, PostgreSQL, and version control systems, are open-source or freely "
        "available for development purposes. Optional expenses are limited to hosted infrastructure and third-party services "
        "such as Render, Neon, Google Maps, Firebase Cloud Messaging, Cloudflare R2-compatible storage, and Gemini API usage."
    ), style="Normal")
    set_paragraph(doc, 226, (
        "The system shall implement automated machine-learning processing for incident categorization, duplicate detection, "
        "risk scoring, toxicity review, media evidence judgment, and analytics or area-insight generation when the configured "
        "provider supports those capabilities."
    ))
    set_paragraph(doc, 230, (
        "The system shall provide a moderator dashboard for reviewing, validating, rejecting, merging, escalating, and "
        "commenting on reports, with additional analytics, user-management, settings, law-enforcement, and admin workflows."
    ))
    set_paragraph(doc, 232, "The system shall provide users with map-based visualization, community feed browsing, saved-area insights, and tracking of reported incidents.")
    set_paragraph(doc, 240, "The use of specific technologies including Expo/React Native, React/Vite, Node.js/Express, PostgreSQL with PostGIS, Socket.IO, FastAPI, Redis-compatible caching, and third-party APIs as the primary development stack.")
    set_paragraph(doc, 273, "CHAPTER III: EXISTING SOLUTIONS")
    set_paragraph(doc, 286, "Established reporting platform patterns")
    set_paragraph(doc, 288, "Enterprise-scale processing pattern")
    set_paragraph(doc, 292, "Institutional design approach")
    set_paragraph(doc, 295, "Community-centered reporting patterns")
    set_paragraph(doc, 293, (
        "From a design perspective, large-scale reporting platforms emphasize reliable data flow, secure user authentication, "
        "responsive user interaction, and efficient information management. System design is carefully structured to support "
        "real-time reporting, media evidence handling, incident tracking, moderation workflows, and scalable data processing "
        "while maintaining usability and accessibility for a broad range of users."
    ), style="Normal")
    set_paragraph(doc, 296, (
        "Community-driven and independent reporting platforms represent a more accessible category of existing solutions, "
        "particularly relevant to academic projects. These systems often rely on open-source or widely available technologies "
        "such as cross-platform mobile frameworks, web dashboards, REST APIs, cloud-hosted databases, and Python-based "
        "machine-learning or analytics services."
    ), style="Normal")
    set_paragraph(doc, 297, "Open-source development pattern")
    set_paragraph(doc, 300, "Community-centered design approach")

    # Architecture and implementation chapters.
    set_paragraph(doc, 325, (
        "The proposed solution consists of a mobile incident-reporting application developed with Expo and React Native, a "
        "web-based moderator dashboard developed with React and Vite, backend API services developed using Node.js and "
        "Express, a Python FastAPI machine-learning service, and a PostgreSQL/PostGIS database. The implementation separates "
        "the mobile client, dashboard client, backend services, realtime communication, data management, and machine-learning "
        "modules so that each subsystem can be tested and extended independently."
    ))
    set_paragraph(doc, 328, "At the project level, the system architecture is organized into five main layers:")
    set_paragraph(doc, 329, "• Mobile Application Layer: Includes authentication, incident reporting, drafts, offline queueing, media upload, geolocation tagging, community feed, notifications, saved areas, witness prompts, and account preferences.")
    set_paragraph(doc, 331, "• Backend and Realtime Services Layer: Manages REST API communication, user authentication, validation, business logic, Socket.IO events, push-notification coordination, scheduled jobs, and communication between clients, the database, and ML services.")
    set_paragraph(doc, 332, "• Machine Learning and Processing Layer: Handles classification, duplicate detection, risk scoring, toxicity review, advisory media evidence judgment, area insights, dashboard insights, and constellation synthesis through either local models or Gemini-backed provider logic.")
    set_paragraph(doc, 333, "• Database and Data Management Layer: Uses PostgreSQL with PostGIS to store users, incidents, reports, media metadata, moderation actions, ML outputs, comments, notifications, saved areas, follows, disclosure settings, and witness-constellation data.")
    set_paragraph(doc, 337, "• Mobile Application Module: Manages registration, login, email verification, Google sign-in, report creation, draft handling, offline queueing, media uploads, location selection, community feed browsing, maps, saved areas, notifications, witness prompts, and user preferences.")
    set_paragraph(doc, 338, "• Backend API Module: Handles request processing, authentication, role-based authorization, validation, API routing, data persistence, upload handling, scheduled maintenance, notification dispatch, and communication with the ML service.")
    set_paragraph(doc, 339, "• Machine Learning Module: Provides incident categorization, duplicate detection, pairwise duplicate comparison, confidence scoring, toxicity analysis, risk scoring, media evidence judgment, area insights, dashboard insights, and witness-constellation synthesis. The service supports a local provider and a Gemini-backed provider depending on deployment configuration.")
    set_paragraph(doc, 340, "• Moderator Dashboard Module: Provides moderators and administrators with tools for reviewing submitted incidents, validating or rejecting reports, linking duplicates, reviewing ML and media judgment outputs, monitoring analytics, managing users and settings, and administering database or access-request workflows.")
    set_paragraph(doc, 341, "• Real-Time Communication Module: Handles Socket.IO rooms, live notifications, incident timeline updates, notification inbox state, push-token workflows, witness prompt delivery, and synchronization between users, staff roles, and backend services.")
    set_paragraph(doc, 342, "• Database Management Module: Oversees secure storage, retrieval, and synchronization of user data, incidents, reports, comments, notifications, saved areas, media metadata, geospatial fields, moderation history, and system-generated analytical results.")
    set_paragraph(doc, 349, "User Interaction Block: Implements user-facing features including authentication, incident reporting forms, draft management, anonymous reporting options, media uploads, notification management, saved areas, witness prompts, location selection, and map/feed browsing.")
    set_paragraph(doc, 351, "Reporting Processing Block: Manages incident report submission and processing. It handles user-submitted descriptions, photo and video evidence, location metadata, draft state, idempotency keys, offline queueing, validation, storage, classification, risk analysis, media judgment, and duplicate detection.")
    set_paragraph(doc, 354, "Geolocation and Mapping Block: Captures geographic coordinates during report submission, links incidents to precise database locations, supports PostGIS-backed queries, and enables map-based visualization, nearby-incident browsing, saved-area insights, and law-enforcement operations maps.")
    set_paragraph(doc, 356, "Machine Learning Analysis Block: Introduces automated analysis for incident categorization, duplicate detection, pairwise duplicate review, confidence scoring, toxicity review, risk scoring, media evidence judgment, and insight generation. These outputs support human decision making rather than replacing moderator or law-enforcement review.")
    set_paragraph(doc, 358, "Moderation and Resolution Block: Provides functionality for reviewing, validating, rejecting, linking duplicates, activating witness constellations, reviewing timeline comments, monitoring analytics, controlling disclosure settings, and coordinating law-enforcement status transitions throughout the incident lifecycle.")
    set_paragraph(doc, 383, "Input block")
    set_paragraph(doc, 385, "Processing block")
    set_paragraph(doc, 391, "Communication block")
    set_paragraph(doc, 395, (
        "The software implementation constitutes the core of the project. Development was carried out using modern "
        "cross-platform mobile development technologies, a React/Vite dashboard, backend API services, machine-learning "
        "services, realtime communication, and a relational database implemented with PostgreSQL/PostGIS. These technologies "
        "provide modular architecture, efficient API communication, and scalable system development suitable for rapid "
        "prototyping and iterative development."
    ))
    set_paragraph(doc, 400, (
        "The project relies on PostgreSQL with PostGIS support to manage persistent storage of system data. The database "
        "stores user accounts, incidents, reports, moderation queue records, report actions, duplicate links, ML outputs, "
        "incident comments, moderator settings, notifications, seen marks, incident follows, saved areas, and witness "
        "constellation data introduced through migrations."
    ))
    set_paragraph(doc, 401, "The database design supports geospatial location queries, report lifecycle timestamps, disclosure controls, idempotent incident creation, media references, text embeddings, ML verdict metadata, and indexes for dashboard, map, feed, notification, and moderation workflows.")
    set_paragraph(doc, 404, "The main software functionalities implemented in the project include:", style="Normal")
    set_paragraph(doc, 405, "• Incident Reporting System: Allows users to submit safety or crime-related incidents through title and description fields, category and severity data, date/time selection, location tagging, photo and video evidence, optional anonymous submission, draft saving, and offline queueing.", style="Normal")
    set_paragraph(doc, 406, "• Machine Learning Processing System: Enables automated analysis through categorization, duplicate detection, risk scoring, toxicity review, media evidence judgment, area insights, dashboard insights, and constellation synthesis. The ML service can run local models or a Gemini-backed provider depending on configuration.", style="Normal")
    set_paragraph(doc, 407, "• Moderation, Administration, and Law-Enforcement System: Provides dashboard workflows for report review, category updates, duplicate linking or dismissal, rejection, escalation, media judgment retry, user management, access-request administration, settings, law-enforcement queue handling, disclosure settings, and case closure.", style="Normal")
    set_paragraph(doc, 408, "• Data and Report Management System: Supports storage, retrieval, updating, and synchronization of incidents and reports with metadata such as report status, timestamps, location information, media references, ML outputs, comments, follows, notifications, and saved areas.", style="Normal")
    set_paragraph(doc, 409, "• Realtime and Notification System: Handles Socket.IO updates, notification inbox records, Firebase/Expo/Notifee notification flows, weekly digest scheduling, push-token management, and witness prompt delivery.", style="Normal")
    set_paragraph(doc, 410, "These functions are implemented across focused modules rather than a single monolithic component, improving clarity and reducing coupling.", style="Normal")

    # Chapter VI: keep CSV evidence, but align interpretation with current codebase.
    set_paragraph(doc, 442, (
        "This chapter presents the experimental evaluation of SafeSignal, with emphasis on the moderator dashboard and the "
        "workflows that support incident review, access control, analytics, and law-enforcement coordination. A consistency "
        "audit was also performed against the current repository state to distinguish historical testcase observations from "
        "functionality that is now implemented in source code and regression tests."
    ))
    set_paragraph(doc, 443, (
        "The consolidated testcase file contains 62 recorded result rows. Of these, 39 passed and 23 failed, producing an "
        "observed pass rate of 62.9 percent for the evaluated dashboard snapshot. These counts are consistent with the "
        "current `testcases/unified-test-cases.csv` file and provide a historical QA baseline rather than a complete statement "
        "of every current code path."
    ))
    set_paragraph(doc, 445, (
        "The evaluated prototype in the testcase evidence is the SafeSignal web-based moderator dashboard. The current project "
        "state also includes the Expo/React Native mobile application, the Node.js/Express backend, PostgreSQL/PostGIS data "
        "model, Socket.IO realtime layer, and Python FastAPI ML service, so the dashboard results should be interpreted as one "
        "major validation slice of a larger implemented system."
    ))
    set_paragraph(doc, 450, (
        "The experiment followed a functional system-testing approach. Each testcase defined a view, feature, identifier, test "
        "title, description, ordered steps, expected result, actual result, outcome, and bug identifier when applicable. The "
        "repository audit also reviewed implemented routes, screens, service modules, database schema, package manifests, and "
        "automated test files to verify whether the report descriptions matched the present implementation."
    ))
    set_paragraph(doc, 452, (
        "The principal metrics were the number of recorded result rows, pass and fail counts, observed pass rate, number of "
        "unique bug identifiers, and failure distribution by functional area. The current repository also contains automated "
        "backend and service tests covering authentication, privacy, role access, report idempotency, moderation state "
        "transitions, media judgment, constellation workflows, statistics, and map behavior."
    ))
    set_paragraph(doc, 465, (
        "A limitation recorded in the original dashboard testcase snapshot was that some restricted-route attempts redirected "
        "without clear feedback. The current dashboard source now routes forbidden access with explicit access-denied feedback, "
        "so this issue should be treated as a historical finding that requires retesting rather than as an unverified current "
        "defect."
    ))
    set_paragraph(doc, 467, (
        "The reports workflow showed a mixed historical result. Keyboard shortcuts and the queue/detail splitter behaved as "
        "expected, while some CSV observations reported data-source inconsistencies and status-transition defects. The current "
        "backend includes validation and regression coverage for selected moderation transitions, including preventing an "
        "already rejected report from being rejected again; however, cross-view report-count consistency still requires "
        "end-to-end retesting against a synchronized dataset."
    ))
    set_paragraph(doc, 468, (
        "These historical results indicate that the moderation interface contains useful interaction mechanics while also "
        "showing why state validation and list refresh behavior are important. Current source-level safeguards address some "
        "of the recorded transition risks, but the dashboard should still be retested end to end after any status-changing "
        "operation."
    ))
    set_paragraph(doc, 473, (
        "User administration tests found both functional and inconsistent behavior in the historical dataset. The current "
        "backend test suite includes coverage for pending staff-user visibility and role/privacy behavior, so access-request "
        "state should be retested in the running dashboard before being reported as an open issue."
    ))
    set_paragraph(doc, 474, (
        "Settings tests identified interaction and feedback issues in the historical dashboard run, including controls that "
        "were difficult to operate or insufficiently labeled. The current repository contains settings APIs and dashboard "
        "settings screens, but a fresh accessibility and interaction pass is still needed to confirm whether each browser, "
        "sound, profile, password, and preference control behaves consistently."
    ))
    set_paragraph(doc, 479, (
        "The overall testcase result remains 39 passing rows and 23 failing rows out of 62 recorded rows. Authentication, "
        "logged-out routing, keyboard shortcuts, panel resizing, selected analytics period filters, and several baseline "
        "dashboard flows behaved correctly in the recorded snapshot. The most important historical failure clusters were "
        "data consistency across dashboard, reports, and analytics views; report status-transition handling; user and "
        "access-request state consistency; settings control reachability; map resource loading; and role-specific "
        "law-enforcement navigation. Several of these areas now have source-level safeguards or automated regression tests, "
        "so the remaining open risk is the absence of a fresh full end-to-end dashboard retest."
    ))
    set_paragraph(doc, 481, (
        "The results show that SafeSignal is not merely a static interface: multiple end-to-end workflows can be executed, "
        "including login, route protection, report navigation, analytics filtering, and selected law-enforcement detail "
        "interactions. The current codebase also shows implemented support for mobile reporting, offline queueing, notifications, "
        "media judgment, witness constellations, privacy controls, and backend regression tests. Remaining evaluation work should "
        "focus on proving synchronized behavior across the full deployed stack."
    ))
    set_paragraph(doc, 482, (
        "The recorded data-consistency failures are significant because SafeSignal depends on trustworthy incident counts and "
        "report status. If dashboard totals, report queues, and analytics cards do not agree, moderators may make decisions from "
        "conflicting information. Historical report-state and user-state failures also motivated stricter backend safeguards and "
        "should be included in the next integrated regression pass."
    ))
    set_paragraph(doc, 485, (
        "The main limitation of this evaluation is that the recorded testcase evidence was performed on a local prototype "
        "dataset and local dashboard server. The test files do not provide load testing, security penetration testing, mobile "
        "application field testing, or machine-learning precision and recall measurements. External Google Maps resources also "
        "failed in the test environment because of proxy-related resource errors, so map behavior should be rechecked in a "
        "network environment where the map service is reachable."
    ))
    set_paragraph(doc, 491, (
        "This chapter presented the experimental evaluation of SafeSignal using evidence from the project testcase records and "
        "dashboard artifacts, and aligned that evidence with the current repository state. The evaluation covered 62 recorded "
        "testcase results across authentication, routing, dashboard analytics, reports, users, settings, law-enforcement workflows, "
        "and runtime behavior. The result was 39 passes and 23 failures, corresponding to an observed pass rate of 62.9 percent."
    ))
    set_paragraph(doc, 492, (
        "The experiments confirm that SafeSignal has a working functional foundation, especially in authentication, basic "
        "navigation, selected analytics filters, report-interface interactions, and role-protected workflows. The current codebase "
        "adds evidence of progress through backend and service tests for moderation state transitions, privacy, notifications, "
        "constellations, and media judgment. The next validation cycle should run a fresh integrated dashboard/mobile/backend/ML "
        "test pass against one synchronized dataset."
    ))

    # Conclusion and future work.
    set_paragraph(doc, 497, (
        "The project successfully delivered a functional mobile-based application and supporting dashboard ecosystem integrating "
        "core features such as incident reporting, media and location submission, draft and offline handling, notifications, "
        "machine-learning-based incident classification, duplicate detection, risk analysis, media evidence judgment, witness "
        "corroboration, moderator verification workflows, law-enforcement coordination, analytics, and secure backend data "
        "management."
    ))
    set_paragraph(doc, 499, (
        "From a technical perspective, the project applied key software engineering principles including modular design, iterative "
        "development, machine-learning integration, role-based access control, privacy-aware data handling, realtime communication, "
        "and continuous testing. The recorded experiments validated an important dashboard slice of the prototype, while the current "
        "source code and automated tests provide additional evidence for backend, ML, privacy, and workflow behavior."
    ))
    set_paragraph(doc, 511, (
        "While the current implementation fulfills the objectives of a functional prototype, several avenues for future work and "
        "improvement have been identified. Since realtime notifications, incident timelines, law-enforcement coordination, saved "
        "areas, and witness prompts already exist at prototype level, future work should focus on hardening these features through "
        "end-to-end retesting, clearer accessibility feedback, synchronized dashboard data sources, production monitoring, and "
        "more complete deployment validation."
    ))
    set_paragraph(doc, 513, (
        "From a technical standpoint, future iterations could improve machine-learning evaluation with precision, recall, and "
        "false-positive analysis for classification, duplicate detection, media judgment, and risk scoring. Additional improvements "
        "include load testing, security testing, durable cloud media storage, stricter privacy review, stronger accessibility testing, "
        "and full Render/Neon/Firebase/Gemini production-readiness checks."
    ))

    # References and end matter placeholders.
    set_paragraph(doc, 558, "No formally captioned figures are included in this version of the report. The cover-page image is a decorative/branding asset and is not treated as a numbered technical figure.")
    set_paragraph(doc, 561, "")
    set_paragraph(doc, 562, "")
    set_paragraph(doc, 568, "The following tables are used in the report:")
    set_paragraph(doc, 570, "Table 1: Student and project information")
    set_paragraph(doc, 571, "Table 2: Project timeline and planning")

    # Appendix: project structure and implementation excerpts.
    set_paragraph(doc, 627, "• Mobile-part/: Expo/React Native mobile app, navigation, screens, hooks, notification services, API clients, offline queueing, and assets.")
    set_paragraph(doc, 628, "• backend/: Node.js/Express API, routes, middleware, services, database schema, migrations, Socket.IO server setup, jobs, tests, and upload handling.")
    set_paragraph(doc, 629, "• moderator-dashboard/: React/Vite dashboard, pages, layouts, reusable components, API service modules, route protection, law-enforcement workspace, settings, analytics, and admin tools.")
    set_paragraph(doc, 630, "• ml-service/: Python FastAPI service, provider abstraction, local and Gemini providers, cache manager, models, utilities, scripts, Dockerfile, and ML tests.")
    set_paragraph(doc, 631, "• constants/: Shared category, status, limit, spacing, typography, theme, and preference constants reused across project components.")
    set_paragraph(doc, 635, (
        "The incident reporting system is implemented through the mobile reporting screen and supporting hooks responsible for "
        "handling form state, validation, draft persistence, media selection, location picking, ML preview toggles, anonymous "
        "defaults, offline queueing, and communication with backend incident APIs."
    ))
    set_paragraph(doc, 637, "• User authentication and identity verification • Incident title, description, photo, and video submission • Location tagging using current location or map selection • Draft saving and offline queueing • Optional anonymous submission • ML-assisted category and risk preview • Structured API communication with idempotency support")
    set_paragraph(doc, 641, (
        "The machine-learning processing system is designed as an independent FastAPI service interacting with backend processing "
        "modules. Submitted reports can be analyzed for categorization, risk, toxicity, duplicate candidates, contextual duplicate "
        "comparison, media evidence judgment, area insights, dashboard insights, and witness-constellation synthesis depending on "
        "the configured provider."
    ))
    set_paragraph(doc, 643, "• Incident categorization based on submitted report content • Duplicate detection through embeddings and optional pairwise comparison • Risk scoring and toxicity review • Advisory media evidence judgment for photos and videos • Analytics and area insight generation • Redis or in-memory cache support for repeated ML calls")
    set_paragraph(doc, 647, (
        "The moderation system is implemented through a role-protected dashboard that allows staff to review incidents, inspect "
        "ML and media evidence outputs, verify or reject reports, update categories, link or dismiss duplicates, activate witness "
        "constellations, communicate through timelines, and coordinate law-enforcement follow-up."
    ))
    set_paragraph(doc, 649, "• Reviewing submitted reports • Approving, rejecting, escalating, or updating incidents • Managing duplicate report verification • Retrying media judgment • Viewing timelines and internal comments • Managing users, settings, access requests, analytics, and law-enforcement case status")
    set_paragraph(doc, 653, (
        "Backend services are responsible for API communication, request validation, data processing, user authentication, role "
        "authorization, upload handling, realtime coordination, notification delivery, scheduled jobs, and system-wide coordination "
        "between application modules. Persistent data storage is managed through PostgreSQL with PostGIS geospatial support."
    ))
    set_paragraph(doc, 655, "• Secure storage of user, incident, report, media, and ML information • Geospatial querying through PostGIS • Notifications, follows, saved areas, comments, and witness constellation data • Reliable communication between mobile, dashboard, backend, and ML services • Consistent synchronization of report and moderation data")

    # Table 2 had a blank task description for ML work.
    doc.tables[1].cell(7, 1).text = "Machine learning service development"

    # Insertions are intentionally delayed until after index-based replacements.
    inserted = insert_after(functional_requirements_anchor, "The system shall support role-based access control for citizens, moderators, administrators, and law-enforcement users.", style="Normal")
    insert_after(inserted, "The system shall support incident drafts, offline report queueing, notification inbox management, and role-aware privacy controls such as anonymous reporting, location fuzzing, and media disclosure settings.", style="Normal")

    insert_after(architecture_layer_anchor, "• Moderator, Admin, and Law-Enforcement Dashboard Layer: Provides report review, analytics, user administration, settings, database administration, law-enforcement incident handling, disclosure controls, and timeline messaging.", style="Normal")

    deployment_heading = insert_after(deployment_body_anchor, "Deployment and hosting configuration", style="Heading3")
    deployment_para = insert_after(deployment_heading, (
        "The project includes a production-oriented deployment path based on Render, Neon, and Cloudflare R2. Render is used "
        "as the target hosting platform for the Node.js/Express backend and the Python FastAPI machine-learning service. The "
        "backend connects to the hosted ML service through the `ML_SERVICE_URL` environment variable, while each service keeps "
        "its own build and runtime configuration."
    ), style="Normal")
    deployment_para = insert_after(deployment_para, (
        "Neon is used as the hosted PostgreSQL database provider. The backend reads the Neon connection string through "
        "`DATABASE_URL`, and the schema relies on PostgreSQL with PostGIS enabled for geospatial incident data. After the "
        "database is created, the project initialization scripts create the application tables, indexes, moderation records, "
        "notification tables, saved areas, and witness-constellation structures."
    ), style="Normal")
    deployment_para = insert_after(deployment_para, (
        "Cloudflare R2 is used for durable media storage when R2 credentials are configured. The upload middleware switches "
        "from local disk storage to R2 when `R2_ACCESS_KEY_ID` and `R2_SECRET_ACCESS_KEY` are present, then stores uploaded "
        "incident photos and videos in the configured R2 bucket using an S3-compatible client. In local development or when "
        "R2 credentials are absent, the backend falls back to local upload storage."
    ), style="Normal")
    insert_after(deployment_para, (
        "The mobile deployed profile points the APK to the hosted backend URL, so a rebuilt mobile application can submit "
        "reports, receive status updates, and communicate with the Render-hosted backend instead of relying on localhost. "
        "This deployment setup keeps application hosting, database persistence, and media storage separated, which improves "
        "maintainability and avoids depending on ephemeral service disk for user evidence."
    ), style="Normal")

    insert_after(risk_heading_anchor, (
        "Risk considerations included data reliability, system scalability, machine-learning accuracy, secure handling of "
        "sensitive user information, location privacy, reliable realtime communication, and prevention of inaccurate or "
        "malicious submissions. The project mitigates these risks through role-based access control, validation, privacy-aware "
        "disclosure settings, moderation workflows, advisory ML outputs, and staged testing. Production-grade load testing, "
        "security testing, mobile field testing, and formal ML accuracy evaluation remain future validation work."
    ), style="Normal")

    inserted = insert_after(references_anchor, "Expo Documentation, Expo Official Documentation. Available at: Expo Documentation", style="Normal")
    inserted = insert_after(inserted, "Firebase Cloud Messaging Documentation, Firebase Documentation. Available at: Firebase Cloud Messaging", style="Normal")
    inserted = insert_after(inserted, "Google Gemini API Documentation, Google AI for Developers. Available at: Gemini API Documentation", style="Normal")
    inserted = insert_after(inserted, "Render Documentation, Render Official Documentation. Available at: Render Documentation", style="Normal")
    inserted = insert_after(inserted, "Neon Documentation, Neon Official Documentation. Available at: Neon Documentation", style="Normal")
    insert_after(inserted, "Cloudflare R2 Documentation, Cloudflare Developers. Available at: Cloudflare R2 Documentation", style="Normal")

    for entry in [
        "Table 3: Comparative analysis of existing platform categories",
        "Table 6.1: Application areas evaluated in the experiment",
        "Table 6.2: Evidence sources used for Chapter VI",
        "Table 6.3: Overall testcase result summary",
        "Table 6.4: Result distribution by functional area",
        "Table 6.5: Representative testcase outcomes",
    ]:
        table_list_anchor = insert_after(table_list_anchor, entry, style=table_list_anchor.style)
        clone_run_format(doc.paragraphs[571], table_list_anchor)

    insert_after(appendix_structure_anchor, "• docs/ and testcases/: Technical notes, deployment guides, feature inventories, and CSV-based QA evidence used to validate and explain the project.", style=appendix_structure_anchor.style)

    # Remove bogus TOC lines caused by body paragraphs styled as headings in the source document.
    for paragraph in toc_paragraphs_to_delete:
        delete_paragraph(paragraph)

    deployment_toc = insert_after(toc_deployment_anchor, "3.3. Deployment and hosting configuration 21", style=toc_deployment_anchor.style)
    clone_run_format(toc_deployment_anchor, deployment_toc)

    acronym_entries = [
        ("Acronym", "Meaning"),
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("APK", "Android Package Kit"),
        ("CI/CD", "Continuous Integration / Continuous Deployment"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("DBMS", "Database Management System"),
        ("FCM", "Firebase Cloud Messaging"),
        ("IEEE", "Institute of Electrical and Electronics Engineers"),
        ("JWT", "JSON Web Token"),
        ("ML", "Machine Learning"),
        ("NLP", "Natural Language Processing"),
        ("ORM", "Object Relational Mapping"),
        ("OTP", "One-Time Password"),
        ("R2", "Cloudflare R2 object storage"),
        ("REST", "Representational State Transfer"),
        ("SDLC", "Software Development Life Cycle"),
        ("SMS", "Short Message Service"),
        ("SQL", "Structured Query Language"),
        ("URL", "Uniform Resource Locator"),
    ]

    for table in list(doc.tables[8:19]):
        delete_table(table)

    acronyms_table = insert_table_after(acronyms_anchor, len(acronym_entries), 2)
    for row_index, (acronym, meaning) in enumerate(acronym_entries):
        acronyms_table.cell(row_index, 0).text = acronym
        acronyms_table.cell(row_index, 1).text = meaning

    normalize_heading_styles(doc)
    for table in doc.tables:
        if len(table.rows) > 1:
            mark_row_as_table_header(table.rows[0])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
