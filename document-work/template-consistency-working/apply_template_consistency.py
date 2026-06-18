from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "document-work" / "FYP Report_Rabih hajj hassan_Version_2.docx"
OUTPUT = ROOT / "document-work" / "FYP Report_Rabih hajj hassan_Version_2_template_consistent.docx"


def mark_row_as_table_header(row):
    tr_props = row._tr.get_or_add_trPr()
    if tr_props.find(qn("w:tblHeader")) is not None:
        return
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_props.append(header)


def set_style_if_present(paragraph, style_name):
    try:
        paragraph.style = style_name
    except KeyError:
        return


def normalize_heading_styles(doc):
    style_names = [style.name for style in doc.styles]
    if "Heading 41" in style_names and "Heading 4" not in style_names:
        doc.styles["Heading 41"].name = "Heading 4"

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 41":
            set_style_if_present(paragraph, "Heading 4")
        if paragraph.style.name == "isselectedend":
            set_style_if_present(paragraph, "Normal")
        if paragraph.style.name.startswith("Heading") and not text:
            set_style_if_present(paragraph, "Normal")


def ensure_num_pr(paragraph, ilvl, num_id):
    p_props = paragraph._p.get_or_add_pPr()
    num_props = p_props.find(qn("w:numPr"))
    if num_props is None:
        num_props = OxmlElement("w:numPr")
        p_props.append(num_props)

    ilvl_el = num_props.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        num_props.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))

    num_id_el = num_props.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_props.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def patch_heading4_style(docx_path):
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()

    with ZipFile(docx_path, "r") as zin, ZipFile(tmp.name, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "word/styles.xml":
                root = etree.fromstring(data)
                style = root.xpath('.//w:style[@w:styleId="Heading41"]', namespaces=ns)
                if style:
                    style = style[0]
                    style.set(f"{{{w_ns}}}styleId", "Heading4")
                    name = style.find("w:name", namespaces=ns)
                    if name is None:
                        name = etree.SubElement(style, f"{{{w_ns}}}name")
                    name.set(f"{{{w_ns}}}val", "Heading 4")

                    p_props = style.find("w:pPr", namespaces=ns)
                    if p_props is None:
                        p_props = etree.SubElement(style, f"{{{w_ns}}}pPr")
                    outline = p_props.find("w:outlineLvl", namespaces=ns)
                    if outline is None:
                        outline = etree.SubElement(p_props, f"{{{w_ns}}}outlineLvl")
                    outline.set(f"{{{w_ns}}}val", "3")

                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

            elif item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    zout.writestr(item, data)
                    continue
                changed = False
                for element in root.xpath('.//*[@w:val="Heading41"]', namespaces=ns):
                    element.set(f"{{{w_ns}}}val", "Heading4")
                    changed = True
                if changed:
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

            zout.writestr(item, data)

    Path(tmp.name).replace(docx_path)


def main():
    doc = Document(SOURCE)

    normalize_heading_styles(doc)

    for index in [53, 55, 57, 59, 61]:
        set_style_if_present(doc.paragraphs[index], "Normal")

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in {"Test result", "Test interpretation", "Discussion"}:
            set_style_if_present(paragraph, "Heading 4")
        if text == "Code Structure and Implementation Excerpts":
            set_style_if_present(paragraph, "Normal")
        if text == "Risk Management and Scope Control":
            ensure_num_pr(paragraph, ilvl=0, num_id=6)
        if text == "Deployment and hosting configuration":
            ensure_num_pr(paragraph, ilvl=1, num_id=19)

    for table in doc.tables:
        if len(table.rows) > 1:
            mark_row_as_table_header(table.rows[0])

    doc.save(OUTPUT)
    patch_heading4_style(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
